from fastapi import APIRouter, HTTPException, Depends, Response
from sqlalchemy import func
from sqlalchemy.orm import Session
from uuid import UUID
import json
import difflib
from .database import get_db
from . import models, schemas
from .workspaces import get_project_in_workspace
from backend.rbac import ensure_project_access
from backend.knowledge_base_service import ensure_workspace_kb, get_relevant_entries
from backend.ai_providers import get_openai_client
from backend.template_service import get_template_version
from backend.ai_guardrails import DECLINE_PHRASE, bundle_context_entries, render_context_block, verify_citations
from backend.prd_service import (
    next_prd_version,
    refresh_prd_embeddings,
    create_decision_embedding,
    search_prd_embeddings,
    build_prd_context_items,
)
from backend.workspace_memory import remember_workspace_event

from fastapi.responses import FileResponse
from docx import Document as DocxDocument
import tempfile
from dotenv import load_dotenv

# 🔑 Load environment variables
load_dotenv()

router = APIRouter(
    prefix="/projects",
    tags=["prds"]
)
embeddings_router = APIRouter(
    prefix="/embeddings",
    tags=["prd-embeddings"],
)


def _context_payload(entries: list[models.KnowledgeBaseEntry]) -> tuple[list[schemas.KnowledgeBaseContextItem], str, set[str]]:
    bundle = bundle_context_entries(entries)
    context_items = [item.to_schema() for item in bundle]
    context_block = render_context_block(bundle)
    allowed_markers = {item.marker for item in context_items if item.marker}
    return context_items, context_block, allowed_markers


def _record_prd_entry(db: Session, workspace_id: UUID | None, prd: models.PRD, user_id: UUID | None) -> None:
    if not workspace_id:
        return
    kb = ensure_workspace_kb(db, workspace_id)
    entry = models.KnowledgeBaseEntry(
        kb_id=kb.id,
        type="prd",
        title=prd.feature_name or prd.project.title if prd.project else "PRD",
        content=prd.content,
        created_by=user_id,
        project_id=prd.project_id,
        tags=["prd"],
    )
    db.add(entry)
    db.commit()
    snippet = (prd.content or "")[:600]
    remember_workspace_event(
        db,
        workspace_id,
        content=f"PRD update for {prd.feature_name or 'PRD'} v{prd.version}:\n{snippet}",
        source="prd",
        metadata={
            "project_id": str(prd.project_id),
            "prd_id": str(prd.id),
            "version": prd.version,
        },
        tags=["prd"],
        user_id=user_id,
    )


def _render_prd_context_block(items: list[schemas.KnowledgeBaseContextItem]) -> str:
    if not items:
        return "No PRD context provided."
    lines: list[str] = []
    for item in items:
        marker = item.marker or ""
        prefix = f"[{marker}] " if marker else ""
        lines.append(f"{prefix}{item.title} -> {item.snippet}")
    return "\n".join(lines)


def _side_by_side_diff(text_a: str | None, text_b: str | None) -> list[schemas.PRDDiffLine]:
    left_lines = (text_a or "").splitlines()
    right_lines = (text_b or "").splitlines()
    matcher = difflib.SequenceMatcher(None, left_lines, right_lines)
    diff: list[schemas.PRDDiffLine] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for offset in range(i2 - i1):
                diff.append(
                    schemas.PRDDiffLine(
                        type="equal",
                        left_line=left_lines[i1 + offset],
                        right_line=right_lines[j1 + offset],
                        left_number=i1 + offset + 1,
                        right_number=j1 + offset + 1,
                    )
                )
        elif tag == "replace":
            span = max(i2 - i1, j2 - j1)
            for offset in range(span):
                left_value = left_lines[i1 + offset] if i1 + offset < i2 else ""
                right_value = right_lines[j1 + offset] if j1 + offset < j2 else ""
                diff.append(
                    schemas.PRDDiffLine(
                        type="replace",
                        left_line=left_value or None,
                        right_line=right_value or None,
                        left_number=i1 + offset + 1 if i1 + offset < i2 else None,
                        right_number=j1 + offset + 1 if j1 + offset < j2 else None,
                    )
                )
        elif tag == "delete":
            for offset, line in enumerate(left_lines[i1:i2]):
                diff.append(
                    schemas.PRDDiffLine(
                        type="delete",
                        left_line=line,
                        left_number=i1 + offset + 1,
                    )
                )
        elif tag == "insert":
            for offset, line in enumerate(right_lines[j1:j2]):
                diff.append(
                    schemas.PRDDiffLine(
                        type="insert",
                        right_line=line,
                        right_number=j1 + offset + 1,
                    )
                )
    return diff

# -----------------------------
# Generate a new PRD (Markdown only)
# -----------------------------
@router.post("/{project_id}/prd", response_model=schemas.PRDResponse)
def generate_prd(
    project_id: str,
    prd_data: schemas.PRDCreate,
    workspace_id: UUID,
    user_id: UUID,
    db: Session = Depends(get_db),
):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="contributor")
    project = get_project_in_workspace(db, project_id, workspace_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    context_query = "\n".join(
        filter(
            None,
            [
                project.title,
                project.description,
                project.goals,
                project.north_star_metric,
                prd_data.feature_name,
                prd_data.prompt,
            ],
        )
    )
    kb_entries = get_relevant_entries(db, project.workspace_id, context_query, top_n=5)
    context_items, context_block, allowed_markers = _context_payload(kb_entries)

    template_section = ""
    if getattr(prd_data, "template_id", None):
        try:
            template, version = get_template_version(db, workspace_id, prd_data.template_id)
            template_section = (
                f"\nUse the following template structure titled '{template.title}':\n{version.content}\n"
            )
        except HTTPException:
            template_section = ""

    # build prompt
    prompt = f"""
Generate a Product Requirements Document (PRD) in **Markdown** format.

Always use Markdown headers, lists, and bullet points.
Include these sections:
- Objective
- Scope
- Success Metrics
- Engineering Requirements
- Future Work

Project Title: {project.title}
Description: {project.description}
Goals: {project.goals}
North Star Metric: {project.north_star_metric or 'Not specified'}
Target Personas: {', '.join(project.target_personas or []) or 'Not specified'}

Workspace Knowledge Base:
{context_block}

Guidance:
- Cite workspace references inline using their markers like [CTX1].
- If the knowledge base does not provide relevant evidence, respond with the exact phrase "{DECLINE_PHRASE}" instead of drafting unsupported content.

Task: Generate a PRD for feature: {prd_data.feature_name}
User instructions: {prd_data.prompt}
{template_section}
"""

    client = get_openai_client(db, workspace_id)
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an expert product manager who writes PRDs."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.4,
    )

    # ✅ FIX: use .content instead of ["content"]
    raw_content = response.choices[0].message.content
    verification = verify_citations([raw_content], allowed_markers)
    if allowed_markers and verification.status == "failed":
        raw_content = DECLINE_PHRASE

    new_prd = models.PRD(
        project_id=project_id,
        workspace_id=project.workspace_id,
        feature_name=prd_data.feature_name,
        description=prd_data.prompt,
        goals=project.goals,
        content=raw_content,
        version=1,
        is_active=True,
        created_by=user_id,
    )
    db.add(new_prd)
    db.commit()
    db.refresh(new_prd)

    _record_prd_entry(db, project.workspace_id, new_prd, user_id)
    refresh_prd_embeddings(db, new_prd)
    db.commit()

    response = schemas.PRDResponse.model_validate(new_prd)
    response.context_entries = context_items
    response.verification = verification
    return response


# ---------------------------------------------------------
# ✅ Refine an existing PRD
# ---------------------------------------------------------
@router.put("/{project_id}/prds/{prd_id}/refine", response_model=schemas.PRDResponse)
def refine_prd(
    project_id: str,
    prd_id: UUID,
    refine_data: schemas.PRDRefine,
    workspace_id: UUID,
    user_id: UUID,
    db: Session = Depends(get_db),
):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="contributor")
    # Fetch the current PRD
    project = get_project_in_workspace(db, project_id, workspace_id)

    prd = (
        db.query(models.PRD)
        .filter(
            models.PRD.id == prd_id,
            models.PRD.project_id == project_id,
            models.PRD.workspace_id == project.workspace_id,
        )
        .first()
    )
    if not prd:
        raise HTTPException(status_code=404, detail="PRD not found")

    # Collect project details
    # project already resolved

    # Collect previous PRDs (excluding current one)
    previous_prds = (
        db.query(models.PRD)
        .filter(
            models.PRD.project_id == project_id,
            models.PRD.id != prd_id,
            models.PRD.workspace_id.in_([project.workspace_id, None]),
        )
        .all()
    )
    previous_prds_text = "\n---\n".join([p.content for p in previous_prds]) if previous_prds else "None"

    refine_query = "\n".join(
        filter(
            None,
            [
                project.title,
                prd.feature_name,
                refine_data.instructions,
                prd.content,
            ],
        )
    )
    kb_entries = get_relevant_entries(db, project.workspace_id, refine_query, top_n=5)
    context_items, context_block, allowed_markers = _context_payload(kb_entries)

    # Build structured context for OpenAI
    messages = [
        {"role": "system", "content": "You are an expert product manager who refines PRDs in clean Markdown format."},
        {"role": "user", "content": f"Project context:\nTitle: {project.title}\nDescription: {project.description}\nGoals: {project.goals}\nNorth Star Metric: {project.north_star_metric or 'Not specified'}"},
        {"role": "user", "content": f"Target Personas: {', '.join(project.target_personas or []) or 'Not specified'}"},
        {"role": "user", "content": f"Workspace knowledge base:\n{context_block}\n\nReference knowledge entries with citations like [CTX1]. If none apply, reply with the exact phrase \"{DECLINE_PHRASE}\" instead of inventing details."},
        {"role": "user", "content": f"Other PRDs for reference:\n{previous_prds_text}"},
        {"role": "user", "content": f"Here is the current PRD to refine:\n{prd.content or ''}"},
        {"role": "user", "content": f"Refinement instructions:\n{refine_data.instructions}\n\nPlease update the PRD accordingly while preserving its structure, goals, and context."},
    ]

    # Call OpenAI
    client = get_openai_client(db, workspace_id)
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,
        temperature=0.2,
        max_tokens=2000,
    )

    refined_content = response.choices[0].message.content
    verification = verify_citations([refined_content], allowed_markers)
    if allowed_markers and verification.status == "failed":
        refined_content = DECLINE_PHRASE

    # Save as a new PRD version
    new_version = next_prd_version(db, project_id, project.workspace_id)
    refined_prd = models.PRD(
        project_id=project_id,
        workspace_id=project.workspace_id,
        feature_name=prd.feature_name,
        description=prd.description,
        goals=prd.goals,
        content=refined_content,
        version=new_version,
        is_active=True,
        created_by=user_id,
    )

    # Mark old PRD inactive
    prd.is_active = False

    db.add(refined_prd)
    db.commit()
    db.refresh(refined_prd)
    _record_prd_entry(db, project.workspace_id, refined_prd, user_id)
    refresh_prd_embeddings(db, refined_prd)
    db.commit()
    response = schemas.PRDResponse.model_validate(refined_prd)
    response.context_entries = context_items
    response.verification = verification
    return response


@router.post("/{project_id}/prds/{prd_id}/save", response_model=schemas.PRDResponse)
def save_prd_version(
    project_id: str,
    prd_id: UUID,
    payload: schemas.PRDSaveRequest,
    workspace_id: UUID,
    user_id: UUID,
    db: Session = Depends(get_db),
):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="contributor")
    project = get_project_in_workspace(db, project_id, workspace_id)
    base_prd = (
        db.query(models.PRD)
        .filter(
            models.PRD.id == prd_id,
            models.PRD.project_id == project_id,
            models.PRD.workspace_id == project.workspace_id,
        )
        .first()
    )
    if not base_prd:
        raise HTTPException(status_code=404, detail="PRD not found")
    content = (payload.content or "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="Content cannot be empty.")

    max_version = next_prd_version(db, project_id, project.workspace_id)
    base_prd.is_active = False
    new_prd = models.PRD(
        project_id=project_id,
        workspace_id=project.workspace_id,
        feature_name=payload.feature_name or base_prd.feature_name,
        description=payload.description or base_prd.description,
        goals=base_prd.goals,
        content=content,
        version=max_version,
        is_active=True,
        created_by=user_id,
    )
    db.add(new_prd)
    db.commit()
    db.refresh(new_prd)
    _record_prd_entry(db, project.workspace_id, new_prd, user_id)
    refresh_prd_embeddings(db, new_prd)
    db.commit()
    return schemas.PRDResponse.model_validate(new_prd)


# -----------------------------
# List all PRDs for a project
# -----------------------------
@router.get("/{project_id}/prds", response_model=list[schemas.PRDResponse])
def list_prds(project_id: str, workspace_id: UUID, user_id: UUID, db: Session = Depends(get_db)):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="viewer")
    project = get_project_in_workspace(db, project_id, workspace_id)
    return (
        db.query(models.PRD)
        .filter(models.PRD.project_id == project_id, models.PRD.workspace_id == project.workspace_id)
        .order_by(models.PRD.version.desc())
        .all()
    )


@router.get("/{project_id}/prds/history", response_model=list[schemas.PRDVersionSummary])
def get_prd_history(project_id: str, workspace_id: UUID, user_id: UUID, db: Session = Depends(get_db)):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="viewer")
    project = get_project_in_workspace(db, project_id, workspace_id)
    rows = (
        db.query(
            models.PRD,
            models.User.display_name.label("author_name"),
            func.count(models.PRDDecisionNote.id).label("decision_count"),
        )
        .outerjoin(models.User, models.PRD.created_by == models.User.id)
        .outerjoin(models.PRDDecisionNote, models.PRDDecisionNote.prd_id == models.PRD.id)
        .filter(models.PRD.project_id == project_id, models.PRD.workspace_id == project.workspace_id)
        .group_by(models.PRD.id, models.User.display_name)
        .order_by(models.PRD.version.desc())
        .all()
    )
    history: list[schemas.PRDVersionSummary] = []
    for prd, author_name, decision_count in rows:
        history.append(
            schemas.PRDVersionSummary(
                id=prd.id,
                version=prd.version,
                feature_name=prd.feature_name,
                is_active=prd.is_active,
                created_at=prd.created_at,
                created_by=prd.created_by,
                author_name=author_name,
                decision_count=int(decision_count or 0),
            )
        )
    return history


# -----------------------------
# Get a specific PRD by ID
# -----------------------------
@router.get("/{project_id}/prds/{prd_id}", response_model=schemas.PRDResponse)
def get_prd(project_id: str, prd_id: UUID, workspace_id: UUID, user_id: UUID, db: Session = Depends(get_db)):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="viewer")
    project = get_project_in_workspace(db, project_id, workspace_id)
    prd = (
        db.query(models.PRD)
        .filter(models.PRD.id == prd_id, models.PRD.project_id == project_id, models.PRD.workspace_id == project.workspace_id)
        .first()
    )
    if not prd:
        raise HTTPException(status_code=404, detail="PRD not found")
    return prd


@router.get("/{project_id}/prds/compare", response_model=schemas.PRDDiffResponse)
def compare_prds(
    project_id: str,
    workspace_id: UUID,
    user_id: UUID,
    v1: int,
    v2: int,
    db: Session = Depends(get_db),
):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="viewer")
    if v1 == v2:
        raise HTTPException(status_code=400, detail="Select two different versions to compare.")
    project = get_project_in_workspace(db, project_id, workspace_id)
    versions = {v1, v2}
    records = (
        db.query(models.PRD)
        .filter(
            models.PRD.project_id == project_id,
            models.PRD.workspace_id == project.workspace_id,
            models.PRD.version.in_(versions),
        )
        .all()
    )
    if len(records) < 2:
        raise HTTPException(status_code=404, detail="Requested versions were not found.")
    version_map = {record.version: record for record in records}
    if v1 not in version_map or v2 not in version_map:
        raise HTTPException(status_code=404, detail="Requested versions were not found.")
    diff = _side_by_side_diff(version_map[v1].content, version_map[v2].content)
    return schemas.PRDDiffResponse(
        version_a=v1,
        version_b=v2,
        prd_a_id=version_map[v1].id,
        prd_b_id=version_map[v2].id,
        diff=diff,
    )


@router.get("/{project_id}/prds/{prd_id}/decisions", response_model=list[schemas.PRDDecisionNoteResponse])
def list_prd_decisions(
    project_id: str,
    prd_id: UUID,
    workspace_id: UUID,
    user_id: UUID,
    version: int | None = None,
    db: Session = Depends(get_db),
):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="viewer")
    project = get_project_in_workspace(db, project_id, workspace_id)
    prd = (
        db.query(models.PRD)
        .filter(
            models.PRD.id == prd_id,
            models.PRD.project_id == project_id,
            models.PRD.workspace_id == project.workspace_id,
        )
        .first()
    )
    if not prd:
        raise HTTPException(status_code=404, detail="PRD not found")
    query = (
        db.query(
            models.PRDDecisionNote,
            models.User.display_name.label("author_name"),
        )
        .outerjoin(models.User, models.PRDDecisionNote.created_by == models.User.id)
        .filter(models.PRDDecisionNote.prd_id == prd_id)
        .order_by(models.PRDDecisionNote.created_at.desc())
    )
    if version is not None:
        query = query.filter(models.PRDDecisionNote.version == version)
    rows = query.all()
    notes: list[schemas.PRDDecisionNoteResponse] = []
    for note, author_name in rows:
        notes.append(
            schemas.PRDDecisionNoteResponse(
                id=note.id,
                prd_id=note.prd_id,
                project_id=note.project_id,
                workspace_id=note.workspace_id,
                version=note.version,
                decision=note.decision,
                rationale=note.rationale,
                created_by=note.created_by,
                author_name=author_name,
                created_at=note.created_at,
            )
        )
    return notes


@router.post("/{project_id}/prds/{prd_id}/decisions", response_model=schemas.PRDDecisionNoteResponse)
def add_prd_decision(
    project_id: str,
    prd_id: UUID,
    payload: schemas.PRDDecisionNoteCreate,
    workspace_id: UUID,
    user_id: UUID,
    db: Session = Depends(get_db),
):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="contributor")
    project = get_project_in_workspace(db, project_id, workspace_id)
    prd = (
        db.query(models.PRD)
        .filter(
            models.PRD.id == prd_id,
            models.PRD.project_id == project_id,
            models.PRD.workspace_id == project.workspace_id,
        )
        .first()
    )
    if not prd:
        raise HTTPException(status_code=404, detail="PRD not found")
    decision = (payload.decision or "").strip()
    if not decision:
        raise HTTPException(status_code=400, detail="Decision cannot be empty.")
    note = models.PRDDecisionNote(
        prd_id=prd.id,
        project_id=project.id,
        workspace_id=project.workspace_id,
        version=payload.version or prd.version,
        decision=decision,
        rationale=(payload.rationale or "").strip() or None,
        created_by=user_id,
    )
    db.add(note)
    db.commit()
    db.refresh(note)
    create_decision_embedding(db, note)
    db.commit()
    author = db.query(models.User).filter(models.User.id == note.created_by).first()
    return schemas.PRDDecisionNoteResponse(
        id=note.id,
        prd_id=note.prd_id,
        project_id=note.project_id,
        workspace_id=note.workspace_id,
        version=note.version,
        decision=note.decision,
        rationale=note.rationale,
        created_by=note.created_by,
        author_name=author.display_name if author else None,
        created_at=note.created_at,
    )


@router.post("/{project_id}/prds/qa", response_model=schemas.PRDQAResponse)
def prd_question_answer(
    project_id: str,
    payload: schemas.PRDQARequest,
    workspace_id: UUID,
    user_id: UUID,
    db: Session = Depends(get_db),
):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="viewer")
    question = (payload.question or "").strip()
    if not question:
        raise HTTPException(status_code=400, detail="Question cannot be empty.")
    project = get_project_in_workspace(db, project_id, workspace_id)
    version_filters: list[int] = []
    if payload.version_a is not None:
        version_filters.append(payload.version_a)
    if payload.version_b is not None and payload.version_b not in version_filters:
        version_filters.append(payload.version_b)
    records = search_prd_embeddings(
        db,
        project.workspace_id,
        UUID(project_id),
        question,
        versions=version_filters or None,
        limit=8,
    )
    context_items = build_prd_context_items(records)
    context_block = _render_prd_context_block(context_items)
    allowed_markers = {item.marker for item in context_items if item.marker}
    used_versions = sorted({record.version for record in records})

    prompt = (
        "You are an AI product partner that answers questions using versioned PRDs and decision notes.\n"
        f"Use the provided snippets only. Cite evidence with markers like [CTX1]. "
        f"If the context is insufficient, reply with the exact phrase \"{DECLINE_PHRASE}\".\n"
        f"Project: {project.title}\n"
        f"Question: {question}\n"
        f"PRD context:\n{context_block}"
    )
    answer = DECLINE_PHRASE if not context_items else "I'm still reviewing prior versions. Try again."
    try:
        client = get_openai_client(db, workspace_id)
        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0.2,
            messages=[
                {"role": "system", "content": "You summarize PRD changes and decisions with precise citations."},
                {"role": "user", "content": prompt},
            ],
        )
        answer = completion.choices[0].message.content or answer
    except Exception:
        pass
    verification = verify_citations([answer], allowed_markers)
    if allowed_markers and verification.status == "failed":
        answer = DECLINE_PHRASE
    response = schemas.PRDQAResponse(
        answer=answer,
        context_entries=context_items,
        used_versions=used_versions,
        verification=verification,
    )
    return response


# -----------------------------
# Delete a PRD
# -----------------------------
@router.delete("/{project_id}/prds/{prd_id}", status_code=204)
def delete_prd(project_id: str, prd_id: UUID, workspace_id: UUID, user_id: UUID, db: Session = Depends(get_db)):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="contributor")
    project = get_project_in_workspace(db, project_id, workspace_id)
    prd = (
        db.query(models.PRD)
        .filter(
            models.PRD.id == prd_id,
            models.PRD.project_id == project_id,
            models.PRD.workspace_id == project.workspace_id,
        )
        .first()
    )

    if not prd:
        raise HTTPException(status_code=404, detail="PRD not found")

    was_active = prd.is_active

    db.delete(prd)
    db.flush()  # ensure deleted record is not considered in follow-up queries

    if was_active:
        replacement = (
            db.query(models.PRD)
            .filter(
                models.PRD.project_id == project_id,
                models.PRD.workspace_id == project.workspace_id,
            )
            .order_by(models.PRD.version.desc(), models.PRD.created_at.desc())
            .first()
        )
        if replacement:
            replacement.is_active = True

    db.commit()

    return Response(status_code=204)


# -----------------------------
# Get the active PRD
# -----------------------------
@router.get("/{project_id}/prd/active", response_model=schemas.PRDResponse)
def get_active_prd(project_id: str, workspace_id: UUID, user_id: UUID, db: Session = Depends(get_db)):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="viewer")
    prd = (
        db.query(models.PRD)
        .filter(
            models.PRD.project_id == project_id,
            models.PRD.workspace_id == workspace_id,
            models.PRD.is_active == True,
        )
        .order_by(models.PRD.version.desc())
        .first()
    )
    if not prd:
        raise HTTPException(status_code=404, detail="No active PRD found")
    return prd


# -----------------------------
# Export PRD to .docx
# -----------------------------
@router.get("/{project_id}/prds/{prd_id}/export")
def export_prd(project_id: str, prd_id: UUID, workspace_id: UUID, user_id: UUID, db: Session = Depends(get_db)):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="viewer")
    prd = (
        db.query(models.PRD)
        .filter(
            models.PRD.id == prd_id,
            models.PRD.project_id == project_id,
            models.PRD.workspace_id == workspace_id,
        )
        .first()
    )
    if not prd:
        raise HTTPException(status_code=404, detail="PRD not found")

    # Create Word document
    doc = DocxDocument()
    doc.add_heading("Product Requirements Document (PRD)", level=0)

    # Add sections
    doc.add_heading("Feature Name", level=1)
    doc.add_paragraph(prd.feature_name or "")

    doc.add_heading("Objective", level=1)
    doc.add_paragraph(prd.description or "")

    doc.add_heading("Goals", level=1)
    doc.add_paragraph(prd.goals or "")

    if hasattr(prd, "content") and prd.content:
        doc.add_heading("Full Markdown PRD", level=1)
        doc.add_paragraph(prd.content)

    # Save to temp file
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)

    return FileResponse(
        tmp.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"PRD_{prd.id}.docx"
    )


@embeddings_router.post("/rebuild/{project_id}")
def rebuild_prd_embeddings_endpoint(
    project_id: str,
    workspace_id: UUID,
    user_id: UUID,
    db: Session = Depends(get_db),
):
    ensure_project_access(db, workspace_id, UUID(project_id), user_id, required_role="contributor")
    project = get_project_in_workspace(db, project_id, workspace_id)
    prds = (
        db.query(models.PRD)
        .filter(models.PRD.project_id == project_id, models.PRD.workspace_id == project.workspace_id)
        .order_by(models.PRD.version.asc())
        .all()
    )
    indexed = 0
    for prd in prds:
        refresh_prd_embeddings(db, prd)
        indexed += 1
    db.commit()
    return {"project_id": project_id, "indexed_versions": indexed}
